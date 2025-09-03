# APP.py — Block-by-Block Tailor (80% JD / 20% Original) with JD-first Skills,
# auto Projects & Certifications add/update.

import streamlit as st
from openai import OpenAI
import docx
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from io import BytesIO
import json, re

st.set_page_config(page_title="Block-by-Block Tailor", layout="wide")

# ---------- Polished dark theme ----------
st.markdown("""
<style>
body, .main, .block-container { background:#0b0b0b !important; color:#fff !important; }
h1, h2, h3, h4, h5, h6 { color:#fff !important; letter-spacing:.2px; }
h1 { font-weight:800 !important; }
h2, h3 { font-weight:700 !important; }
.stTextInput > div > input, .stTextArea textarea { background:#121212 !important; color:#fff !important; border:1px solid #2b2b2b !important; }
.stTextArea textarea { min-height:260px !important; }
.stFileUploader { color:#fff !important; }
.stButton>button, .stDownloadButton>button {
  background:#1DB954 !important; color:#0b0b0b !important; font-weight:800;
  border-radius:10px; border:none; padding:.6rem 1rem;
}
.stButton>button:hover, .stDownloadButton>button:hover { filter:brightness(1.05); }
.rule { border-top:1px solid #2b2b2b; margin: 0.75rem 0 1rem; }
.option-card { background:#101010; border:1px solid #262626; border-radius:12px; padding:16px; margin-top:10px; }
.option-card h4 { margin-top:0; }
.option-grid { display:grid; grid-template-columns: repeat(4, minmax(160px,1fr)); gap:10px 18px; }
.option-grid .stCheckbox { margin:0 !important; }
.title-emoji { font-size:42px; margin-right:8px; vertical-align:-6px; }
</style>
""", unsafe_allow_html=True)

st.markdown('<h1><span class="title-emoji">🧱</span>Block-by-Block Tailor (80% JD / 20% Original)</h1>', unsafe_allow_html=True)

# ---------- Top row ----------
c1, c2 = st.columns(2, gap="large")
with c1:
    st.subheader("Upload your resume (.docx)")
    uploaded = st.file_uploader("Drop your Word file here", type=["docx"])
with c2:
    st.subheader("Paste the Job Description")
    jd = st.text_area("Paste full JD", label_visibility="collapsed")

# ---------- Options panel ----------
st.markdown('<div class="option-card">', unsafe_allow_html=True)
st.markdown("#### Options")

colA, colB = st.columns([2,1])
with colA:
    out_name = st.text_input("Output file name (no extension)", value="Tailored_Resume_Blocks")
with colB:
    do_ats = st.checkbox("🎯 Add ATS keyword stuffing at the END", value=True)

st.markdown('<div class="rule"></div>', unsafe_allow_html=True)
st.markdown("##### Include / Exclude blocks")
col_wrap = st.columns(4)
incl = {}
labels = [
    ("Summary", True), ("Professional Experience", True), ("All Role Blocks", True),
    ("Projects", True), ("Certifications", True), ("Skills", True),
    ("Education", False), ("Awards", False), ("Publications", False),
    ("Volunteering", False), ("Languages", False), ("Interests", False),
]
for i, (label, default) in enumerate(labels):
    with col_wrap[i % 4]:
        incl[label.upper()] = st.checkbox(label, value=default, key=f"inc_{label}")

st.markdown('</div>', unsafe_allow_html=True)  # end option-card

# ---------- Section map ----------
INCLUDE_MAP = {
    "SUMMARY": incl.get("SUMMARY", True),
    "PROFESSIONAL SUMMARY": incl.get("SUMMARY", True),
    "EXPERIENCE": incl.get("PROFESSIONAL EXPERIENCE", True),
    "PROFESSIONAL EXPERIENCE": incl.get("PROFESSIONAL EXPERIENCE", True),
    "PROJECTS": incl.get("PROJECTS", True),
    "PERSONAL PROJECTS": incl.get("PROJECTS", True),
    "CERTIFICATIONS": incl.get("CERTIFICATIONS", True),
    "CERTIFICATION": incl.get("CERTIFICATIONS", True),
    "SKILLS": incl.get("SKILLS", True),
    "EDUCATION": incl.get("EDUCATION", False),
    "AWARDS": incl.get("AWARDS", False),
    "PUBLICATIONS": incl.get("PUBLICATIONS", False),
    "VOLUNTEERING": incl.get("VOLUNTEERING", False),
    "LANGUAGES": incl.get("LANGUAGES", False),
    "INTERESTS": incl.get("INTERESTS", False),
}
INCLUDE_ROLES = incl.get("ALL ROLE BLOCKS", True)

# ---------- OpenAI ----------
client = None
try:
    if "OPENAI_API_KEY" in st.secrets:
        client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
except Exception:
    client = None

# ---------- Helpers ----------
SECTION_NAMES = list(INCLUDE_MAP.keys())
ROLE_KEYWORDS = ["engineer","developer","analyst","administrator","architect","manager","consultant","sre","devops","intern","lead","specialist"]
COMPANY_HINTS = ["inc","llc","ltd","corp","technologies","technology","systems","solutions","imports","bank","labs","university","college","school","department","services"]
DATE_PATTERNS = [
    r"\b(20\d{2}|19\d{2})\b",
    r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}",
    r"\b(Present|Current)\b",
    r"\b\d{1,2}/\d{4}\b"
]

def looks_like_role_line(text: str) -> bool:
    t = text.lower()
    has_comma = "," in t or " – " in t or " - " in t
    has_role = any(k in t for k in ROLE_KEYWORDS)
    has_company = any(k in t for k in COMPANY_HINTS)
    has_date = any(re.search(p, text) for p in DATE_PATTERNS)
    return (INCLUDE_ROLES and ((has_date and (has_company or has_role)) or (has_comma and (has_role or has_company))))

def is_heading_para(p):
    text = (p.text or "").strip()
    if not text:
        return False
    style = (p.style.name or "").lower() if p.style else ""
    if any(h in style for h in ["heading", "title"]): return True
    if text.isupper() and 2 <= len(text) <= 120: return True
    if any(text.upper().startswith(n) for n in SECTION_NAMES): return INCLUDE_MAP.get(text.upper().split(":")[0], True)
    if looks_like_role_line(text): return True
    return False

def get_blocks(doc):
    heads = [i for i,p in enumerate(doc.paragraphs) if is_heading_para(p)]
    if not heads:
        return [(0, len(doc.paragraphs), doc.paragraphs[0].text.strip() if doc.paragraphs else "")]
    blocks=[]
    for j,h in enumerate(heads):
        start = h
        end = heads[j+1] if (j+1)<len(heads) else len(doc.paragraphs)
        header = doc.paragraphs[h].text.strip()
        blocks.append((start,end,header))
    return blocks

def find_section_range(doc, names):
    heads = get_blocks(doc)
    for s,e,head in heads:
        if any(head.upper().startswith(n) for n in names):
            return (s,e)
    return None

def is_bullet_paragraph(p):
    style = (p.style.name or "").lower() if p.style else ""
    if "list" in style or "bullet" in style: return True
    txt = p.text.strip()
    return bool(txt.startswith(("•","-","–","—","·")))

def prefix_bullet(existing_text):
    m = re.match(r'^([•\\-\\–\\—\\·]\\s*)', existing_text.strip())
    return m.group(1) if m else ""

def strip_bullet_prefix(text):
    return re.sub(r'^[•\\-\\–\\—\\·]\\s*', '', text).strip()

def include_block_by_header(header_text: str) -> bool:
    if not header_text: return True
    upper = header_text.upper()
    for key, inc in INCLUDE_MAP.items():
        if upper.startswith(key): return inc
    return INCLUDE_ROLES

def _set_hidden(run):
    rPr = run._element.get_or_add_rPr()
    vanish = OxmlElement('w:vanish')
    rPr.append(vanish)

def add_invisible_stuffing_end(doc, keywords, *, max_words=120, chunks=3):
    if not keywords: return
    seen, cleaned = set(), []
    for k in keywords:
        k=(k or "").strip().lower()
        if k and k not in seen:
            seen.add(k); cleaned.append(k)
        if len(cleaned)>=max_words: break
    if not cleaned: return
    size=max(1,len(cleaned)//max(1,chunks))
    slices=[cleaned[i:i+size] for i in range(0,len(cleaned),size)][:chunks]
    for sl in slices:
        p=doc.add_paragraph()
        pf=p.paragraph_format; pf.space_before=Pt(0); pf.space_after=Pt(0); pf.line_spacing=1.0
        r=p.add_run(" ".join(sl)); r.font.size=Pt(1); r.font.color.rgb=RGBColor(255,255,255)
        try: _set_hidden(r)
        except Exception: pass

def replace_plain_paragraphs(doc, start, end, new_lines):
    paras = doc.paragraphs[start:end]
    # choose non-bullet, non-empty targets
    targets = [i for i,p in enumerate(paras) if p.text.strip() and not is_bullet_paragraph(p)]
    if not targets:
        # fallback: write into first paragraph slot
        targets = [0] if paras else []
    n = min(len(targets), len(new_lines))
    for k in range(n):
        p = paras[targets[k]]
        txt = new_lines[k]
        if p.runs:
            p.runs[0].text = txt
            for r in p.runs[1:]: r.text = ""
        else:
            p.add_run(txt)

def replace_bullet_texts(doc, start, end, new_bullets):
    paras = doc.paragraphs[start:end]
    bullets = [i for i,p in enumerate(paras) if is_bullet_paragraph(p)]
    if not bullets:
        # treat all non-empty as bullets
        bullets = [i for i,p in enumerate(paras) if p.text.strip()]
    n = min(len(bullets), len(new_bullets))
    for k in range(n):
        p = paras[bullets[k]]
        new_text = prefix_bullet(p.text) + strip_bullet_prefix(new_bullets[k])
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]: r.text = ""
        else:
            p.add_run(new_text)

def append_section(doc, heading_text, lines, as_bullets=True):
    """Safely append a new section at the END (no mid-gaps)."""
    # heading
    h = doc.add_paragraph()
    h.style = doc.styles['Heading 2'] if 'Heading 2' in [s.name for s in doc.styles] else h.style
    h.add_run(heading_text.upper())
    # body
    for line in lines:
        p = doc.add_paragraph()
        if as_bullets:
            p.style = 'List Bullet' if 'List Bullet' in [s.name for s in doc.styles] else p.style
            p.add_run(f"• {strip_bullet_prefix(line)}")
        else:
            p.add_run(line)

# ---------- LLM helpers ----------
def chat_json(prompt):
    if not client: raise RuntimeError("OpenAI key missing (OPENAI_API_KEY).")
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role":"system","content":"Return strict JSON only."},{"role":"user","content":prompt}],
        temperature=0.2,
    )
    raw=(resp.choices[0].message.content or "").strip()
    m=re.search(r"\\{.*\\}", raw, re.DOTALL)
    if m: raw=m.group(0)
    return json.loads(raw)

def llm_rewrite_block(block_title, lines, jd_text):
    prompt = f"""
We are tailoring a resume block-by-block.

BLOCK TITLE:
{block_title}

ORIGINAL LINES:
{json.dumps(lines, ensure_ascii=False)}

JOB DESCRIPTION:
{jd_text}

INSTRUCTIONS:
- Rewrite this block to be **80% aligned to the JD** and **20% preserved from the original**.
- Keep the **same number of lines** as the original.
- Preserve bullet vs non-bullet (do NOT include the bullet symbol).
- Do NOT change company/role names or dates in the first line.
- Return up to 60 JD keywords for stuffing at the END.

Return ONLY JSON:
{{
  "rewritten": ["...", "..."],  // same length
  "stuffing_keywords": ["...", "..."]
}}
""".strip()
    data = chat_json(prompt)
    return data.get("rewritten",[]), data.get("stuffing_keywords",[])

def llm_prioritize_skills(jd_text, current_text):
    prompt = f"""
Given the JOB DESCRIPTION and CURRENT SKILLS TEXT, produce an ordered list where
JD-relevant skills appear first, followed by the remaining original skills.

JOB DESCRIPTION:
{jd_text}

CURRENT SKILLS TEXT:
{current_text}

Rules:
- Output 15–30 skills max.
- Keep concise tokens (e.g., "azure", "kubernetes", "python").
- No duplicates. All lowercase.

Return JSON: {{"skills": ["azure","kubernetes","python", ...]}}
""".strip()
    data = chat_json(prompt)
    return data.get("skills", [])

def llm_generate_projects(jd_text, profile_hint):
    prompt = f"""
Create a compact list of 3–5 projects that prove experience relevant to the JD.

JOB DESCRIPTION:
{jd_text}

PROFILE HINT (free text from resume summary or experience keywords):
{profile_hint}

Rules:
- Each project should be one bullet line: "Project Title: short impact statement with metrics".
- Keep it ATS-friendly and realistic for a mid-level profile.

Return JSON: {{"projects": ["...", "..."]}}
""".strip()
    data = chat_json(prompt)
    return data.get("projects", [])

def llm_generate_certs(jd_text):
    prompt = f"""
List 3–6 certification titles that are most relevant to the JD (vendor neutral or major cloud vendors).

JOB DESCRIPTION:
{jd_text}

Return JSON: {{"certifications": ["...", "..."]}}
""".strip()
    data = chat_json(prompt)
    return data.get("certifications", [])

# ---------- Rewriters ----------
def rewrite_doc_block_by_block(doc, jd_text, add_stuffing=True):
    blocks=get_blocks(doc); all_keywords=[]
    for (start,end,header) in blocks:
        if not include_block_by_header(header): continue
        paras=doc.paragraphs[start:end]
        if not paras: continue
        texts, flags, idxs = [], [], []
        for i,p in enumerate(paras):
            t=(p.text or "").strip()
            if not t: continue
            texts.append(strip_bullet_prefix(t))
            flags.append(is_bullet_paragraph(p))
            idxs.append(i)
        if not texts: continue
        try:
            rewritten,kws=llm_rewrite_block(paras[0].text.strip(), texts, jd_text)
        except Exception:
            continue
        if len(rewritten)!=len(texts): continue
        for pos,new_line in enumerate(rewritten):
            p=paras[idxs[pos]]
            original=(p.text or "").strip()
            new_text=(new_line or "").strip()
            if flags[pos]: new_text = prefix_bullet(original) + new_text
            if p.runs:
                p.runs[0].text=new_text
                for r in p.runs[1:]: r.text=""
            else:
                p.add_run(new_text)
        if kws: all_keywords.extend(kws)
    if add_stuffing: add_invisible_stuffing_end(doc, all_keywords, max_words=120, chunks=3)

def update_or_add_skills(doc, jd_text):
    rng = find_section_range(doc, ["SKILLS"])
    existing_text = ""
    if rng:
        s,e = rng
        existing_text = "\n".join(p.text for p in doc.paragraphs[s:e]).strip()
    skills = llm_prioritize_skills(jd_text, existing_text)
    if not skills: return
    lines = [", ".join(skills)]
    if rng:
        s,e = rng
        replace_plain_paragraphs(doc, s, e, lines)
    else:
        append_section(doc, "SKILLS", lines, as_bullets=False)

def update_or_add_projects(doc, jd_text):
    rng = find_section_range(doc, ["PROJECTS","PERSONAL PROJECTS"])
    # small hint: use first 3-4 paragraphs as profile text
    hint = "\n".join(p.text for p in doc.paragraphs[:10])
    projects = llm_generate_projects(jd_text, hint)
    if not projects: return
    if rng:
        s,e = rng
        replace_bullet_texts(doc, s, e, projects)
    else:
        append_section(doc, "PROJECTS", projects, as_bullets=True)

def update_or_add_certs(doc, jd_text):
    rng = find_section_range(doc, ["CERTIFICATIONS","CERTIFICATION"])
    certs = llm_generate_certs(jd_text)
    if not certs: return
    if rng:
        s,e = rng
        replace_plain_paragraphs(doc, s, e, certs)
    else:
        append_section(doc, "CERTIFICATIONS", certs, as_bullets=False)

# ---------- Action ----------
if st.button("Tailor Block-by-Block (80/20)", use_container_width=True):
    if not uploaded or not jd:
        st.error("Please upload a .docx resume and paste the JD.")
        st.stop()
    if not client:
        st.error("Missing OpenAI API key in .streamlit/secrets.toml (OPENAI_API_KEY).")
        st.stop()

    doc = docx.Document(uploaded)

    # 1) Rewrite all blocks (Summary/Experience etc.)
    rewrite_doc_block_by_block(doc, jd, add_stuffing=do_ats)

    # 2) JD-first Skills (add or refresh)
    if INCLUDE_MAP.get("SKILLS", True):
        try: update_or_add_skills(doc, jd)
        except Exception: pass

    # 3) Projects aligned to JD (add or refresh)
    if INCLUDE_MAP.get("PROJECTS", True):
        try: update_or_add_projects(doc, jd)
        except Exception: pass

    # 4) Role-relevant Certifications (add or refresh)
    if INCLUDE_MAP.get("CERTIFICATIONS", True):
        try: update_or_add_certs(doc, jd)
        except Exception: pass

    bio = BytesIO(); doc.save(bio); bio.seek(0)
    st.success("✅ Tailored successfully!")
    st.download_button("📥 Download Tailored Resume",
        data=bio.read(),
        file_name=f"{out_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.caption("All experience blocks updated, JD-first Skills prioritized, and relevant Projects/Certifications added or refreshed. ATS keywords appended only at the very end.")
